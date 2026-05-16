from docx import Document
import os
import pandas as pd
import re
import numpy as np
from dateutil.relativedelta import relativedelta
from pythainlp.util import bahttext
from num2words import num2words
from docx2pdf import convert
from docx.shared import Pt

base_path = os.getcwd()
excel_file = os.path.join(base_path, 'data_input.xlsx')
lease_agreement = os.path.join(base_path, 'lease-contract-public.docx')
destination_folder = os.path.join(base_path, 'output')
owner_contract = os.path.join(base_path, 'owner-contract-public.docx')

LATE_FEE_LOW = 500
LATE_FEE_HIGH = 1000
LATE_FEE_THRESHOLD = 30000

SUFFIX_FONT_NAME = 'Cordia New (Body CS)'
SUFFIX_FONT_SIZE_PT = 15

df = pd.read_excel(excel_file)
df_flipped = df.set_index('Attributes').transpose()

os.makedirs(destination_folder, exist_ok=True)

thai_month_names = {
    1: 'มกราคม',
    2: 'กุมภาพันธ์',
    3: 'มีนาคม',
    4: 'เมษายน',
    5: 'พฤษภาคม',
    6: 'มิถุนายน',
    7: 'กรกฎาคม',
    8: 'สิงหาคม',
    9: 'กันยายน',
    10: 'ตุลาคม',
    11: 'พฤศจิกายน',
    12: 'ธันวาคม'
}


def check_for_thai_name(df, en_col, th_col):
    df[th_col] = df[th_col].fillna(df[en_col])


def replace_text_without_format(paragraph, old_text, new_text):
    new_text = str(new_text) if pd.notna(new_text) else ''
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)


def idorpassport(df, input_column, output_column):
    df[[f'{output_column}_en', f'{output_column}_th']] = df[input_column].apply(lambda x: ('ID card' if isinstance(x, str) and x.lower(
    ) in ['thailand', 'thai'] else 'passport', 'บัตรประชาชน' if isinstance(x, str) and x.lower() in ['thailand', 'thai'] else 'หนังสือเดินทาง')).apply(pd.Series)


def number_to_text(df, column):
    df[column +
        '_text_en'] = df[column].apply(lambda x: num2words(x).replace(", ", " "))
    df[column + '_text_th'] = df[column].apply(bahttext)


def calculate_two_months_deposit(df, column):
    df[column + '_times_two'] = df[column].apply(lambda x: x * 2)


def lease_period(df, start_date, end_date):
    duration_list = []
    duration_list_th = []
    for start, end in zip(df[start_date], df[end_date]):
        end = end + pd.DateOffset(days=1)
        duration = relativedelta(end, start)

        components = []
        components_th = []
        if duration.years:
            components.append(f"{duration.years} Year" +
                              ("s" if duration.years > 1 else ""))
            components_th.append(f'{duration.years} ปี')
        if duration.months:
            components.append(f"{duration.months} Month" +
                              ("s" if duration.months > 1 else ""))
            components_th.append(f'{duration.months} เดือน')
        if duration.days:
            components.append(f"{duration.days} Day" +
                              ("s" if duration.days > 1 else ""))
            components_th.append(f'{duration.days} วัน')

        formatted_duration_th = " ".join(components_th)
        formatted_duration = ", ".join(components)

        duration_list.append(formatted_duration)
        duration_list_th.append(formatted_duration_th)

    df['lease_period_en'] = duration_list
    df['lease_period_th'] = duration_list_th


def late_fee_calculate(df, rent_price_column):
    df['late_fee_num'] = df[rent_price_column].apply(
        lambda x: LATE_FEE_LOW if x < LATE_FEE_THRESHOLD else LATE_FEE_HIGH)


def format_day_with_suffix(day):
    if 10 <= day % 100 <= 20:
        suffix = 'thsuffixplahor'
    else:
        suffix = {1: 'stsuffixplahor', 2: 'ndsuffixplahor',
                  3: 'rdsuffixplahor'}.get(day % 10, 'thsuffixplahor')

    return f"{day}{suffix}"


def convert_date_format(df, column):
    df[column] = df[column].replace(r'^\s*$', np.nan, regex=True)
    df[column] = pd.to_datetime(
        df[column], errors='coerce')
    if df[column].notna().all():
        df[f'{column}_day'] = df[column].dt.day
        df[f'{column}_month'] = df[column].dt.month
        df[f'{column}_year'] = df[column].dt.year
        df[f'{column}_month_en'] = df[column].dt.month_name()
        df[f'{column}_month_th'] = df[f'{column}_month'].map(
            thai_month_names)
        df[column + '_year_th'] = df[column +
                                     '_year'].apply(lambda x: f'พ.ศ. {x + 543}')
        df[column + '_en'] = df[column + '_day'].astype(
            str) + ' ' + df[column + '_month_en'] + ' ' + df[column + '_year'].astype(str)
        df[column + '_th'] = df[column + '_day'].astype(
            str) + ' ' + df[column + '_month_th'] + ' ' + df[column + '_year_th'].astype(str)
        df[column + '_day_en_suffix'] = df[f'{column}_day'].apply(
            lambda x: format_day_with_suffix(x))
        df[column + '_month_year_en'] = df[column + '_month_en'] + \
            ' ' + df[column + '_year'].astype(str)


def format_room(room):
    if re.fullmatch(r"\d+/\d+", room):
        room = re.sub(r"(\d+)/(\d+)", r"\1-\2", room)
        room = f"({room})"
    else:
        room = re.sub(r"(\d+)/(\d+)", r"\1-\2", room)
    return room


def replace_text_with_format(paragraph, old_text, new_text):
    new_text = str(new_text)
    for run in paragraph.runs:
        if old_text in run.text:
            if old_text.isupper():
                new_text = new_text.upper()
            elif old_text.islower():
                new_text = new_text.lower()
            elif old_text.istitle():
                new_text = new_text.title()

            run.text = run.text.replace(old_text, new_text)


def replace_text_in_tables(table, old_text, new_text):
    for table_row in table.rows:
        for cell_row in table_row.cells:
            if cell_row.paragraphs:
                for cell_paragraph in cell_row.paragraphs:
                    replace_text_with_format(
                        cell_paragraph, old_text, new_text if pd.notna(new_text) else '')


def replace_text_if_df_exist(paragraph, old_text, new_text):
    replace_text_with_format(
        paragraph, old_text, new_text if pd.notna(new_text) else '')


def add_comma(df, column):
    df[column] = df[column].apply(lambda x: "{:,}".format(x))


def pay_duration(df, start_day_column):

    conditions = [
        df[start_day_column] <= 26,
        df[start_day_column] == 27,
        df[start_day_column] == 28,
        df[start_day_column] == 29,
        df[start_day_column] == 30,
        df[start_day_column] == 31
    ]

    choices = [
        (df[start_day_column] + 4).apply(format_day_with_suffix),
        "31stsuffixplahor or 1stsuffixplahor",
        "1stsuffixplahor or 2ndsuffixplahor",
        "2ndsuffixplahor or 3rdsuffixplahor",
        "3rdsuffixplahor or 4thsuffixplahor",
        "4thsuffixplahor"
    ]

    choicesth = [
        (df[start_day_column] + 4).astype(str),
        "31 หรือ 1",
        "1 หรือ 2",
        "2 หรือ 3",
        "3 หรือ 4",
        "4"
    ]

    df["last_day_before_fee"] = np.select(
        conditions, choices, default="Date Error")

    df["last_day_before_fee_th"] = np.select(
        conditions, choicesth, default="Date Error")
    return df


def replace_suffix(paragraph, old_text, new_text):
    for run in paragraph.runs:
        if old_text in run.text:
            before, _, after = run.text.partition(old_text)

            run.text = before

            superscript_run = paragraph.add_run(new_text)
            superscript_run.font.superscript = True
            superscript_run.font.name = SUFFIX_FONT_NAME
            superscript_run.font.size = Pt(SUFFIX_FONT_SIZE_PT)

            run.element.addnext(superscript_run.element)

            if after:
                after_run = paragraph.add_run(after)
                after_run.font.name = SUFFIX_FONT_NAME
                after_run.font.size = Pt(SUFFIX_FONT_SIZE_PT)
                superscript_run._element.addnext(after_run._element)


convert_date_format(df_flipped, "rent_start_date")
convert_date_format(df_flipped, "rent_end_date")
convert_date_format(df_flipped, "owner_passport_expire_date")
convert_date_format(df_flipped, "owner_passport_expire_date_2")
convert_date_format(df_flipped, "tenant_passport_expire_date")
convert_date_format(df_flipped, "tenant_passport_expire_date_2")


lease_period(df_flipped, "rent_start_date", "rent_end_date")
calculate_two_months_deposit(df_flipped, "rent_price")
pay_duration(df_flipped, "rent_start_date_day")
late_fee_calculate(df_flipped, 'rent_price')

number_to_text(df_flipped, "rent_price")
number_to_text(df_flipped, "rent_price_times_two")
number_to_text(df_flipped, 'late_fee_num')

add_comma(df_flipped, 'rent_price')
add_comma(df_flipped, 'rent_price_times_two')
add_comma(df_flipped, 'late_fee_num')

idorpassport(df_flipped, 'owner_nationality', 'ow1idp')
idorpassport(df_flipped, 'owner_nationality_2', 'ow2idp')
idorpassport(df_flipped, 'tenant_nationality', 'te1idp')
idorpassport(df_flipped, 'tenant_nationality_2', 'te2idp')

check_for_thai_name(df_flipped, 'owner_name', 'owner_name_th')
check_for_thai_name(df_flipped, 'owner_name_2', 'owner_name_2_th')
check_for_thai_name(df_flipped, 'tenant_name', 'tenant_name_th')
check_for_thai_name(df_flipped, 'tenant_name_2', 'tenant_name_2_th')


df_flipped['room_floor_number_suf'] = df_flipped['room_floor_number'].apply(
    format_day_with_suffix)
df_flipped['owner_name_2_ft'] = df_flipped['owner_name_2'].apply(
    lambda x: f'({x})' if pd.notna(x) else x)
df_flipped['tenant_name_2_ft'] = df_flipped['tenant_name_2'].apply(
    lambda x: f'({x})' if pd.notna(x) else x)
df_flipped['witness_name_2_ft'] = df_flipped['witness_name_2'].apply(
    lambda x: f'({x})' if pd.notna(x) else x)


df_flipped = df_flipped.fillna('')

for index, row in df_flipped.iterrows():
    doc = Document(owner_contract)

    for paragraph in doc.paragraphs:
        replace_text_with_format(
            paragraph, 'PROJECTNAMEHOLDER', row['project_name'])
        replace_text_with_format(
            paragraph, 'UNITNUMBERHOLDER', row['room_number'])
        replace_text_with_format(
            paragraph, 'STYAHD', row['rent_start_date_year'])
        replace_text_with_format(
            paragraph, 'STMDHD', row['rent_start_date_month'])
        replace_text_with_format(
            paragraph, 'STDYHD', row['rent_start_date_day'])
        replace_text_with_format(
            paragraph, 'ENYAHD', row['rent_end_date_year'])
        replace_text_with_format(
            paragraph, 'ENMDHD', row['rent_end_date_month'])
        replace_text_with_format(
            paragraph, 'ENDYHD', row['rent_end_date_day'])

    for table in doc.tables:
        replace_text_in_tables(table, 'NAME1HOLDER', row['owner_name'])
        replace_text_in_tables(table, 'NAME2HOLDER',
                               f"/ {row['owner_name_2']}" if row['owner_name_2'] else '')
        replace_text_in_tables(
            table, 'NAME2NOSLASHHOLDER', row['owner_name_2'])
        replace_text_in_tables(table, 'PSPT1HO', row['owner_passport'])
        replace_text_in_tables(
            table, 'PSPT2HO', f"/ {row['owner_passport_2']}" if row['owner_passport_2'] else '')

    room_num = format_room(str(row['room_number']))

    file_path1 = os.path.join(
        destination_folder, f"代租管合約-{str(row['project_name'])} {room_num}.docx")
    iteration1 = 1
    while os.path.exists(file_path1):
        file_path1 = os.path.join(
            destination_folder, f"代租管合約-{str(row['project_name'])} {room_num}_{str(iteration1)}.docx")
        iteration1 += 1
    doc.save(file_path1)

    pdf_file_path1 = file_path1.replace('.docx', '.pdf')
    iterationpdf1 = 1
    while os.path.exists(pdf_file_path1):
        pdf_file_path1 = os.path.join(
            destination_folder, f"代租管合約-{str(row['project_name'])} {room_num}_{str(iterationpdf1)}.pdf")
        iterationpdf1 += 1
    convert(file_path1, pdf_file_path1)
    print(f"File has been saved to {file_path1}")
    print(f"File has been saved to {pdf_file_path1}")

paragraph_placeholder_dictionary = {
    'startdayplahor': 'rent_start_date_day_en_suffix',
    'Startmmyyplahor': 'rent_start_date_month_year_en',
    'Startdatethplahor': 'rent_start_date_th',
    'enddayplahor': 'rent_end_date_day_en_suffix',
    'Endmmyyplahor': 'rent_end_date_month_year_en',
    'Enddatethplahor': 'rent_end_date_th',
    'OWNERNAMEPLAHOR': 'owner_name',
    'OWNERNAMETHPLAHOR': 'owner_name_th',
    'OWNERPASSPLAHOR': 'owner_passport',
    'Ownernatplahor': 'owner_nationality',
    'Ownernatthplahor': 'owner_nationality_th',
    'Ownerpassexpplahor': 'owner_passport_expire_date_en',
    'Ownerpassexpthplahor': 'owner_passport_expire_date_th',
    'Projectnameplahor': 'project_name',
    'Roomnumberplahor': 'room_number',
    'Buildingnumberplahor': 'building_no',
    'Projectaddressplahor': 'project_address',
    'Projectaddressthplahor': 'project_address_th',
    'TENANTNAMEPLAHOR': 'tenant_name',
    'TENANTNAMETHPLAHOR': 'tenant_name_th',
    'TENANTPASSPLAHOR': 'tenant_passport',
    'Tenantnatplahor': 'tenant_nationality',
    'Tenantnatthplahor': 'tenant_nationality_th',
    'Tenantpassexpplahor': 'tenant_passport_expire_date_en',
    'Tenantpassexpthplahor': 'tenant_passport_expire_date_th',
    'floorplahor': 'room_floor_number',
    'floorsufplahor': 'room_floor_number_suf',
    'Areaplahor': 'room_area',
    'rentstartdatedayplahor': 'rent_start_date_day',
    'lastdaybeforefeeplahor': 'last_day_before_fee',
    'lastdaybeforefeethplahor': 'last_day_before_fee_th',
    'Rentnoplahor': 'rent_price',
    'Rentenplahor': 'rent_price_text_en',
    'Rentthplahor': 'rent_price_text_th',
    'Latefeenumplahor': 'late_fee_num',
    'Latefeetextenplahor': 'late_fee_num_text_en',
    'Latefeetextthlahor': 'late_fee_num_text_th',
    'Leaseperiodenplahor': 'lease_period_en',
    'Leaseperiodthplahor': 'lease_period_th',
    'Depositplahor': 'rent_price_times_two',
    'Deposittextplahor': 'rent_price_times_two_text_en',
    'Deposittextthplahor': 'rent_price_times_two_text_th',
    'OWNBANAPLAHOR': 'owner_bank',
    'OWNBANATHPLAHOR': 'owner_bank_th',
    'Ownbabrplahor': 'owner_bank_branch',
    'Ownbabrthplahor': 'owner_bank_branch_th',
    'Ownbaaccnoplahor': 'owner_bank_account_no',
    'OWNERBANKACCOUNTNAMEPLAHOR': 'owner_bank_account_name',
    'Wameplahor': 'water_meter_no',
    'Elmeplahor': 'electric_meter_no',
}

no_change_format = {
    'ow1idpenplahor': 'ow1idp_en',
    'ow1idpthplahor': 'ow1idp_th',
    'te1idpenplahor': 'te1idp_en',
    'te1idpthplahor': 'te1idp_th',
}

footer_placeholder_dictionary = {
    'OWNERNAMEPLAHOR': 'owner_name',
    'TENANTNAMEPLAHOR': 'tenant_name',
    'WITNESSNAMEPLAHOR': 'witness_name',
    'OWNERNAME2FTPLAHOR': 'owner_name_2_ft',
    'TENANTNAME2FTPLAHOR': 'tenant_name_2_ft',
    'WITNESSNAME2FTPLAHOR': 'witness_name_2_ft',
}
footer_placeholder_sign_dictionary = {

    'Iflandlord2plahor': ('Landlord………………………………', 'owner_name_2'),
    'Iftenant2plahor': ('Tenant…………………………………', 'tenant_name_2'),
    'Ifwitness2plahor': ('Witness………………………………', 'witness_name_2')
}

owner_2_plahor = {
    'Ifowner2dicenplahor': ', OWNERNAME2PLAHOR, holding ow2idpenplahor no. OWNERPASS2PLAHOR of Ownernat2plahor Nationality Expiry date Ownerpassexp2plahor',
    'Ifowner2dicthplahor': ', OWNERNAME2THPLAHOR ถือow2idpthplahorเลขที่ OWNERPASS2PLAHOR สัญชาติ Ownernatth2plahor หมดอายุวันที่ Ownerpassexpth2plahor'

}

tenant_2_plahor = {
    'Iftenant2dicenplahor': ', TENANTNAME2PLAHOR, holding te2idpenplahor no. TENANTPASS2PLAHOR of Tenantnat2plahor Nationality Expiry date Tenantpassexp2plahor',
    'Iftenant2dicthplahor': ', TENANTNAME2THPLAHOR ถือte2idpthplahorเลขที่ TENANTPASS2PLAHOR สัญชาติ Tenantnat2thplahor หมดอายุวันที่ Tenantpassexp2thplahor'

}

owner_2_value_plahor = {
    'OWNERNAME2PLAHOR': 'owner_name_2',
    'OWNERNAME2THPLAHOR': 'owner_name_2_th',
    'OWNERPASS2PLAHOR': 'owner_passport_2',
    'Ownernat2plahor': 'owner_nationality_2',
    'Ownernatth2plahor': 'owner_nationality_2_th',
    'Ownerpassexp2plahor': 'owner_passport_expire_date_2_en',
    'Ownerpassexpth2plahor': 'owner_passport_expire_date_2_th',
    'ow2idpenplahor': 'ow2idp_en',
    'ow2idpthplahor': 'ow2idp_th'

}
tenant_2_value_plahor = {
    'TENANTNAME2PLAHOR': 'tenant_name_2',
    'TENANTNAME2THPLAHOR': 'tenant_name_2_th',
    'TENANTPASS2PLAHOR': 'tenant_passport_2',
    'Tenantnat2plahor': 'tenant_nationality_2',
    'Tenantnat2thplahor': 'tenant_nationality_2_th',
    'Tenantpassexp2plahor': 'tenant_passport_expire_date_2_en',
    'Tenantpassexp2thplahor': 'tenant_passport_expire_date_2_th',
    'te2idpenplahor': 'te2idp_en',
    'te2idpthplahor': 'te2idp_th'
}

superscript_dic = {
    'thsuffixplahor': 'th',
    'stsuffixplahor': 'st',
    'ndsuffixplahor': 'nd',
    'rdsuffixplahor': 'rd',
}

for index2, row2 in df_flipped.iterrows():
    doc2 = Document(lease_agreement)

    for paragraph2 in doc2.paragraphs:
        for placeholder2, column_name2 in paragraph_placeholder_dictionary.items():
            replace_text_if_df_exist(
                paragraph2, placeholder2, row2[column_name2])
        for ncfpla, ncfcol in no_change_format.items():
            replace_text_without_format(paragraph2, ncfpla, row2[ncfcol])

        for owner_2_dic, owner_2_text_plahor in owner_2_plahor.items():
            replace_text_with_format(
                paragraph2, owner_2_dic, owner_2_text_plahor if row2['owner_name_2'] else '')
        for owner_2_plahor_text, owner_2_value_text in owner_2_value_plahor.items():
            replace_text_if_df_exist(
                paragraph2, owner_2_plahor_text, row2[owner_2_value_text] if row2['owner_name_2'] else '')

        for tenant_2_dic, tenant_2_text_plahor in tenant_2_plahor.items():
            replace_text_with_format(
                paragraph2, tenant_2_dic, tenant_2_text_plahor if row2['tenant_name_2'] else '')
        for tenant_2_plahor_text, tenant_2_value_text in tenant_2_value_plahor.items():
            replace_text_if_df_exist(
                paragraph2, tenant_2_plahor_text, row2[tenant_2_value_text] if row2['tenant_name_2'] else '')

        for _ in range(2):
            for suffix_pla, suffix_replace in superscript_dic.items():
                replace_suffix(paragraph2, suffix_pla, suffix_replace)

    for section in doc2.sections:
        footer = section.footer
        for footer_table in footer.tables:
            for footer_holder, footer_column in footer_placeholder_dictionary.items():
                replace_text_in_tables(
                    footer_table, footer_holder, row2[footer_column])
            for footer_sign_holder, footer_sign_replace in footer_placeholder_sign_dictionary.items():
                replace_text_in_tables(footer_table, footer_sign_holder,
                                       footer_sign_replace[0] if row2[footer_sign_replace[1]] else '')

    room_num2 = format_room(str(row2['room_number']))

    file_path2 = os.path.join(
        destination_folder, f"Lease Agreement-{str(row2['project_name'])} {room_num2}.docx")
    iteration2 = 1
    while os.path.exists(file_path2):
        file_path2 = os.path.join(
            destination_folder, f"Lease Agreement-{str(row2['project_name'])} {room_num2}_{str(iteration2)}.docx")
        iteration2 += 1
    doc2.save(file_path2)

    pdf_file_path2 = file_path2.replace('.docx', '.pdf')
    iterationpdf2 = 1
    while os.path.exists(pdf_file_path2):
        pdf_file_path2 = os.path.join(
            destination_folder, f"Lease Agreement-{str(row2['project_name'])} {room_num2}_{str(iterationpdf2)}.pdf")
        iterationpdf2 += 1
    convert(file_path2, pdf_file_path2)

    print(f"File has been saved to {file_path2}")
    print(f"File has been saved to {pdf_file_path2}")
